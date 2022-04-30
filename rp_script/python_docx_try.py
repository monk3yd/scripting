import re
import docx
import pprint
import pandas as pd

def main():
    # Open csv and read all cases
    df = pd.read_csv("BD - Protecciones.csv")[[
        # "TERMINADA",
        "CORTE",
        "RECURRENTE",
        "CI",
        "ISAPRE",
        "RUT ISAPRE",
        "REP ISAPRE",
        "DOMICILIO ISAPRE",
        "FECHA CARTA/FUN",
        "PLAN",
        "% ALZA",
        "PB",
        "PBR",
        "MES OBJECIÓN"
    ]]

    # Clean df --- extract NAPB only
    df_filtered = df.dropna(subset=["FECHA CARTA/FUN"])

    # df to list of dicts, each element is list needs to be a different docs
    df_clean = df_filtered.to_dict(orient="records")

    # Open docx document
    doc_template = docx.Document("MODELO RECURSO ALZA PRECIO BASE 2022.docx")

    # For each client in dataframe create a file
    for client in df_clean:
        client_doc = docx.Document()
        # Check all client's data
        for key, replace_str in client.items():
            print(key, '->', replace_str) 

            regex = re.compile(f"<<{key}>>")
            # regex = re.compile("<<(.*?)>>")
            
            print("regex", "->", regex)
            print("replace_str", "->", replace_str)

            # Iterate through all paragraphs in docs once --> change only one value,key pair(data) per iteration
            for paragraph in doc_template.paragraphs:
                new_paragraph = paragraph_replace_text(paragraph, regex, replace_str)
                pprint.pprint(list(r.text for r in paragraph.runs))
                client_doc.add_paragraph(list(r.text for r in new_paragraph.runs))
            # exit() # Exit when getting first paragraph

        client_doc.save(f"{client['RECURRENTE']}.docx")
        exit() # Exit when creating first client doc

# Thanks to @scanny for this function and library, you can find it at https://github.com/python-openxml/python-docx/issues/30#issuecomment-879593691
def paragraph_replace_text(paragraph, regex, replace_str):
    """Return `paragraph` after replacing all matches for `regex` with `replace_str`.

    `regex` is a compiled regular expression prepared with `re.compile(pattern)`
    according to the Python library documentation for the `re` module.
    """
    # --- a paragraph may contain more than one match, loop until all are replaced ---
    while True:
        text = paragraph.text
        match = regex.search(text)
        if not match:
            break

        # --- when there's a match, we need to modify run.text for each run that
        # --- contains any part of the match-string.
        runs = iter(paragraph.runs)
        start, end = match.start(), match.end()

        # --- Skip over any leading runs that do not contain the match ---
        for run in runs:
            run_len = len(run.text)
            if start < run_len:
                break
            start, end = start - run_len, end - run_len

        # --- Match starts somewhere in the current run. Replace match-str prefix
        # --- occurring in this run with entire replacement str.
        run_text = run.text
        run_len = len(run_text)
        run.text = "%s%s%s" % (run_text[:start], replace_str, run_text[end:])
        end -= run_len  # --- note this is run-len before replacement ---
        break
        # --- Remove any suffix of match word that occurs in following runs. Note that
        # --- such a suffix will always begin at the first character of the run. Also
        # --- note a suffix can span one or more entire following runs.
        # for run in runs:  # --- next and remaining runs, uses same iterator ---
        #     if end <= 0:
        #         break
        #     run_text = run.text
        #     run_len = len(run_text)
        #     run.text = run_text[end:]
        #     end -= run_len

    # --- optionally get rid of any "spanned" runs that are now empty. This
    # --- could potentially delete things like inline pictures, so use your judgement.
    # for run in paragraph.runs:
    #     if run.text == "":
    #         r = run._r
    #         r.getparent().remove(r)
    return paragraph
 


if __name__ == "__main__":
    main()




#     document = Document()
#     paragraph = document.add_paragraph()
#     paragraph.add_run("f").bold = True
#     paragraph.add_run("o").bold = True
#     paragraph.add_run("o to").bold = True
#     paragraph.add_run(" you and ")
#     paragraph.add_run("foo").bold = True
#     paragraph.add_run(" to the horse")
#     paragraph_replace_text(paragraph, regex, "bar")

#     import pprint